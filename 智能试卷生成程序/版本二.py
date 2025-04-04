#导入必要库
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from sqlalchemy import create_engine, Column, Integer, String, JSON, DateTime, Float,text
from sqlalchemy.orm import declarative_base, sessionmaker
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import pyodbc
import os
import urllib.parse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import webbrowser


# ================== SQL Server数据库配置 ==================
# 配置数据库连接
def get_db_connection():
    # 请修改为你的实际数据库配置
    username = "CZQ"
    password = "1357810la"  # 替换为你的密码
    server = "devil"
    database = "exam_db"

    # 对密码中的特殊字符进行编码
    encoded_password = urllib.parse.quote_plus(password)

    # 使用ODBC Driver 17 for SQL Server
    return f"mssql+pyodbc://{username}:{encoded_password}@{server}/{database}?driver=ODBC+Driver+17+for+SQL+Server"


DATABASE_URL = get_db_connection()#获取数据库连接字符串
Base = declarative_base()#创建ORM模型基类


class Question(Base):
    __tablename__ = 'questions'
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(String(500), nullable=False)
    answer = Column(String(500))
    tags = Column(JSON)
    error_count = Column(Integer, default=1)
    difficulty = Column(Float, default=3.0)
    last_tested = Column(DateTime, default=datetime.now)
    created_at = Column(DateTime, default=datetime.now)


# 创建数据库引擎
def create_db_engine():
    try:
        engine = create_engine(
            DATABASE_URL,
            pool_size=5,
            max_overflow=10,
            pool_timeout=30,
            echo=True  # 开启日志查看详细错误
        )
        try:
            with engine.connect() as conn:
                print("尝试执行测试查询...")
                result = conn.execute(text("SELECT 1"))  # 使用text()包装SQL
                print(f"测试查询成功，结果：{result.scalar()}")
            return engine
        except Exception as e:
            print(f"连接测试失败: {str(e)}")
            raise
    except Exception as e:
        #traceback.print_exc()  # 打印完整堆栈跟踪
        messagebox.showerror("数据库错误", f"连接失败: {str(e)}")
        return None



engine = create_db_engine()
if engine is None:
    exit()

Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)


# ================== 主应用程序 ==================
class ExamToolApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("智能错题系统 v3.0 (SQL Server版)")
        self.geometry("1200x700")
        self.protocol("WM_DELETE_WINDOW", self.on_close)

        # 初始化变量
        self.current_paper = []
        self.current_stats = None

        # 创建界面
        self.create_widgets()

        # 加载初始数据
        self.load_questions()
        self.update_stats()

    def create_widgets(self):
        """创建所有界面组件"""
        # 主布局
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        # 标签页1：错题管理
        self.tab1 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text="错题管理")
        self.create_question_tab()

        # 标签页2：统计分析
        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab2, text="统计分析")
        self.create_stats_tab()

        # 标签页3：生成试卷
        self.tab3 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab3, text="生成试卷")
        self.create_paper_tab()

        # 状态栏
        self.status_bar = ttk.Label(self, text="就绪", relief=tk.SUNKEN)
        self.status_bar.pack(fill=tk.X)

    def create_question_tab(self):
        """创建错题管理标签页"""
        # 输入表单框架
        form_frame = ttk.LabelFrame(self.tab1, text="添加/编辑错题")
        form_frame.pack(padx=10, pady=5, fill=tk.X)

        # 题目内容
        ttk.Label(form_frame, text="题目内容:").grid(row=0, column=0, sticky='w', pady=2)
        self.content_entry = tk.Text(form_frame, height=5, width=80, wrap=tk.WORD)
        self.content_entry.grid(row=0, column=1, columnspan=3, padx=5, sticky='ew')

        # 正确答案
        ttk.Label(form_frame, text="正确答案:").grid(row=1, column=0, sticky='w', pady=2)
        self.answer_entry = tk.Text(form_frame, height=3, width=80, wrap=tk.WORD)
        self.answer_entry.grid(row=1, column=1, columnspan=3, padx=5, sticky='ew')

        # 知识点标签
        ttk.Label(form_frame, text="知识点标签:").grid(row=2, column=0, sticky='w', pady=2)
        self.tags_entry = ttk.Entry(form_frame, width=80)
        self.tags_entry.grid(row=2, column=1, columnspan=3, padx=5, sticky='ew')
        ttk.Label(form_frame, text="(多个标签用逗号分隔)").grid(row=3, column=1, sticky='w')

        # 难度和错误次数
        ttk.Label(form_frame, text="难度:").grid(row=4, column=0, sticky='w', pady=2)
        self.difficulty_var = tk.IntVar(value=3)
        ttk.Combobox(form_frame, textvariable=self.difficulty_var,
                     values=[1, 2, 3, 4, 5], width=5).grid(row=4, column=1, sticky='w')

        ttk.Label(form_frame, text="错误次数:").grid(row=4, column=2, sticky='w', padx=10)
        self.error_count_var = tk.IntVar(value=1)
        ttk.Spinbox(form_frame, textvariable=self.error_count_var,
                    from_=1, to=100, width=5).grid(row=4, column=3, sticky='w')

        # 按钮组
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=5, column=1, columnspan=3, pady=10, sticky='e')

        ttk.Button(button_frame, text="清空", command=self.clear_form).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="保存", command=self.save_question).pack(side=tk.LEFT, padx=5)

        # 错题列表框架
        list_frame = ttk.LabelFrame(self.tab1, text="错题列表")
        list_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

        # 列表工具栏
        toolbar = ttk.Frame(list_frame)
        toolbar.pack(fill=tk.X)

        ttk.Button(toolbar, text="刷新", command=self.load_questions).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="标记为已掌握", command=self.mark_as_learned).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="删除选中", command=self.delete_question).pack(side=tk.LEFT, padx=2)

        # 搜索框
        search_frame = ttk.Frame(toolbar)
        search_frame.pack(side=tk.RIGHT, padx=5)

        ttk.Label(search_frame, text="搜索:").pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=20)
        search_entry.pack(side=tk.LEFT, padx=2)
        search_entry.bind("<Return>", lambda e: self.load_questions())
        ttk.Button(search_frame, text="搜索", command=self.load_questions).pack(side=tk.LEFT)

        # 错题列表
        columns = ("id", "content", "tags", "error_count", "difficulty", "last_tested", "created_at")
        self.tree = ttk.Treeview(
            list_frame,
            columns=columns,
            show="headings",
            selectmode="extended"
        )

        # 配置列
        self.tree.column("id", width=50, anchor="center")
        self.tree.column("content", width=300)
        self.tree.column("tags", width=150)
        self.tree.column("error_count", width=80, anchor="center")
        self.tree.column("difficulty", width=80, anchor="center")
        self.tree.column("last_tested", width=120, anchor="center")
        self.tree.column("created_at", width=120, anchor="center")

        # 设置列标题
        self.tree.heading("id", text="ID")
        self.tree.heading("content", text="题目内容")
        self.tree.heading("tags", text="知识点")
        self.tree.heading("error_count", text="错误次数")
        self.tree.heading("difficulty", text="难度")
        self.tree.heading("last_tested", text="最后测试")
        self.tree.heading("created_at", text="创建时间")

        # 添加滚动条
        scroll_y = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        scroll_y.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=scroll_y.set)

        scroll_x = ttk.Scrollbar(list_frame, orient="horizontal", command=self.tree.xview)
        scroll_x.pack(side="bottom", fill="x")
        self.tree.configure(xscrollcommand=scroll_x.set)

        self.tree.pack(fill=tk.BOTH, expand=True)

        # 绑定事件
        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.tree.bind("<Button-3>", self.show_context_menu)

        # 右键菜单
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="编辑", command=self.edit_question)
        self.context_menu.add_command(label="标记为已掌握", command=self.mark_as_learned)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="删除", command=self.delete_question)

    def create_stats_tab(self):
        """创建统计分析标签页"""
        # 统计图框架
        self.stats_frame = ttk.Frame(self.tab2)
        self.stats_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 控制面板
        control_frame = ttk.Frame(self.tab2)
        control_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(control_frame, text="刷新统计", command=self.update_stats).pack(side=tk.LEFT)
        ttk.Button(control_frame, text="导出图表", command=self.export_charts).pack(side=tk.LEFT, padx=5)

        # 时间范围选择
        ttk.Label(control_frame, text="统计时间范围:").pack(side=tk.LEFT, padx=5)
        self.time_range_var = tk.StringVar(value="all")
        ttk.Radiobutton(control_frame, text="全部", variable=self.time_range_var,
                        value="all", command=self.update_stats).pack(side=tk.LEFT)
        ttk.Radiobutton(control_frame, text="最近7天", variable=self.time_range_var,
                        value="7days", command=self.update_stats).pack(side=tk.LEFT)
        ttk.Radiobutton(control_frame, text="最近30天", variable=self.time_range_var,
                        value="30days", command=self.update_stats).pack(side=tk.LEFT)

    def create_paper_tab(self):
        """创建生成试卷标签页"""
        # 设置框架
        settings_frame = ttk.LabelFrame(self.tab3, text="试卷设置")
        settings_frame.pack(padx=10, pady=5, fill=tk.X)

        # 试卷标题
        ttk.Label(settings_frame, text="试卷标题:").grid(row=0, column=0, sticky='w', padx=5, pady=2)
        self.paper_title_var = tk.StringVar(value="错题强化试卷")
        ttk.Entry(settings_frame, textvariable=self.paper_title_var, width=40).grid(row=0, column=1, sticky='ew')

        # 题目数量
        ttk.Label(settings_frame, text="题目数量:").grid(row=1, column=0, sticky='w', padx=5, pady=2)
        self.paper_count_var = tk.IntVar(value=10)
        ttk.Spinbox(settings_frame, textvariable=self.paper_count_var, from_=1, to=100, width=5).grid(row=1, column=1,
                                                                                                      sticky='w')

        # 难度分布
        ttk.Label(settings_frame, text="难度分布:").grid(row=2, column=0, sticky='w', padx=5, pady=2)
        self.difficulty_dist_var = tk.StringVar(value="1:2:3:2:1")
        ttk.Entry(settings_frame, textvariable=self.difficulty_dist_var, width=15).grid(row=2, column=1, sticky='w')
        ttk.Label(settings_frame, text="(格式:简单:较易:中等:较难:困难,如1:2:3:2:1)").grid(row=2, column=2, sticky='w')

        # 知识点筛选
        ttk.Label(settings_frame, text="知识点筛选:").grid(row=3, column=0, sticky='w', padx=5, pady=2)
        self.tags_filter_var = tk.StringVar()
        ttk.Entry(settings_frame, textvariable=self.tags_filter_var, width=40).grid(row=3, column=1, columnspan=2,
                                                                                    sticky='ew')

        # 生成按钮
        button_frame = ttk.Frame(settings_frame)
        button_frame.grid(row=4, column=1, columnspan=2, pady=5, sticky='e')

        ttk.Button(button_frame, text="生成试卷", command=self.generate_paper).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="导出Word", command=self.export_to_word).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="导出PDF", command=self.export_to_pdf).pack(side=tk.LEFT, padx=5)

        # 试卷预览
        preview_frame = ttk.LabelFrame(self.tab3, text="试卷预览")
        preview_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

        self.paper_preview = tk.Text(
            preview_frame,
            wrap=tk.WORD,
            font=("宋体", 10),
            padx=5,
            pady=5
        )

        scroll_y = ttk.Scrollbar(preview_frame, orient="vertical", command=self.paper_preview.yview)
        scroll_y.pack(side="right", fill="y")
        self.paper_preview.configure(yscrollcommand=scroll_y.set)

        scroll_x = ttk.Scrollbar(preview_frame, orient="horizontal", command=self.paper_preview.xview)
        scroll_x.pack(side="bottom", fill="x")
        self.paper_preview.configure(xscrollcommand=scroll_x.set)

        self.paper_preview.pack(fill=tk.BOTH, expand=True)

    # ================== 数据库操作方法 ==================
    def load_questions(self):
        """从数据库加载错题"""
        self.tree.delete(*self.tree.get_children())
        session = Session()
        try:
            query = session.query(Question)

            # 应用搜索条件
            search_text = self.search_var.get().strip()
            if search_text:
                query = query.filter(
                    Question.content.like(f"%{search_text}%") |
                    Question.answer.like(f"%{search_text}%")
                )

            # 按错误次数降序排列
            questions = query.order_by(Question.error_count.desc()).all()

            for q in questions:
                self.tree.insert("", tk.END, values=(
                    q.id,
                    q.content[:100] + "..." if len(q.content) > 100 else q.content,
                    ", ".join(q.tags) if q.tags else "",
                    q.error_count,
                    q.difficulty,
                    q.last_tested.strftime("%Y-%m-%d"),
                    q.created_at.strftime("%Y-%m-%d")
                ))

            self.status_bar.config(text=f"加载完成，共 {len(questions)} 道错题")
        except Exception as e:
            messagebox.showerror("错误", f"加载数据失败:\n{str(e)}")
        finally:
            session.close()

    def save_question(self):
        """保存错题到数据库"""
        content = self.content_entry.get("1.0", tk.END).strip()
        answer = self.answer_entry.get("1.0", tk.END).strip()
        tags = [t.strip() for t in self.tags_entry.get().split(",") if t.strip()]

        if not content:
            messagebox.showwarning("警告", "题目内容不能为空！")
            return

        session = Session()
        try:
            # 检查是否是编辑现有题目
            selected_item = self.tree.focus()
            if selected_item:
                # 编辑模式
                item_id = self.tree.item(selected_item)['values'][0]
                question = session.query(Question).filter_by(id=item_id).first()
                if question:
                    question.content = content
                    question.answer = answer
                    question.tags = tags
                    question.difficulty = self.difficulty_var.get()
                    question.error_count = self.error_count_var.get()
                    question.last_tested = datetime.now()
                    session.commit()
                    messagebox.showinfo("成功", "错题更新成功！")
            else:
                # 新增模式
                new_question = Question(
                    content=content,
                    answer=answer,
                    tags=tags,
                    difficulty=self.difficulty_var.get(),
                    error_count=self.error_count_var.get(),
                    last_tested=datetime.now()
                )
                session.add(new_question)
                session.commit()
                messagebox.showinfo("成功", "新错题添加成功！")

            self.clear_form()
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"保存失败:\n{str(e)}")
        finally:
            session.close()

    def clear_form(self):
        """清空输入表单"""
        self.content_entry.delete("1.0", tk.END)
        self.answer_entry.delete("1.0", tk.END)
        self.tags_entry.delete(0, tk.END)
        self.difficulty_var.set(3)
        self.error_count_var.set(1)
        # 清除树形视图的选择
        for item in self.tree.selection():
            self.tree.selection_remove(item)

    def on_tree_double_click(self, event):
        """双击错题列表项加载到编辑表单"""
        self.edit_question()

    def edit_question(self):
        """编辑选中的错题"""
        selected_item = self.tree.focus()
        if not selected_item:
            return

        item_id = self.tree.item(selected_item)['values'][0]
        session = Session()
        try:
            question = session.query(Question).filter_by(id=item_id).first()
            if question:
                self.clear_form()

                self.content_entry.insert("1.0", question.content)
                self.answer_entry.insert("1.0", question.answer)
                self.tags_entry.insert(0, ", ".join(question.tags) if question.tags else "")
                self.difficulty_var.set(question.difficulty)
                self.error_count_var.set(question.error_count)

                self.status_bar.config(text=f"正在编辑题目 ID: {item_id}")
        except Exception as e:
            messagebox.showerror("错误", f"加载题目失败:\n{str(e)}")
        finally:
            session.close()

    def mark_as_learned(self):
        """标记选中的题目为已掌握"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择题目！")
            return

        if not messagebox.askyesno("确认", f"确定要标记 {len(selected_items)} 道题为已掌握吗？"):
            return

        session = Session()
        try:
            count = 0
            for item in selected_items:
                item_id = self.tree.item(item)['values'][0]
                question = session.query(Question).filter_by(id=item_id).first()
                if question:
                    question.error_count = max(0, question.error_count - 1)
                    question.last_tested = datetime.now()
                    count += 1

            session.commit()
            messagebox.showinfo("成功", f"已标记 {count} 道题为已掌握")
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"标记失败:\n{str(e)}")
        finally:
            session.close()

    def delete_question(self):
        """删除选中的题目"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择题目！")
            return

        if not messagebox.askyesno("确认", f"确定要删除 {len(selected_items)} 道题吗？此操作不可恢复！"):
            return

        session = Session()
        try:
            count = 0
            for item in selected_items:
                item_id = self.tree.item(item)['values'][0]
                question = session.query(Question).filter_by(id=item_id).first()
                if question:
                    session.delete(question)
                    count += 1

            session.commit()
            messagebox.showinfo("成功", f"已删除 {count} 道题")
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"删除失败:\n{str(e)}")
        finally:
            session.close()

    def show_context_menu(self, event):
        """显示右键菜单"""
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    # ================== 统计分析方法 ==================
    def update_stats(self):
        """更新统计图表"""
        # 清除旧内容
        for widget in self.stats_frame.winfo_children():
            widget.destroy()

        session = Session()
        try:
            # 确定时间范围
            time_range = self.time_range_var.get()
            now = datetime.now()

            if time_range == "7days":
                start_date = now - timedelta(days=7)
            elif time_range == "30days":
                start_date = now - timedelta(days=30)
            else:
                start_date = None  # 全部数据

            # 获取数据
            query = session.query(Question)
            if start_date:
                query = query.filter(Question.created_at >= start_date)

            questions = query.all()

            if not questions:
                ttk.Label(self.stats_frame, text="没有找到符合条件的错题数据").pack()
                return

            # 准备统计数据
            stats_data = {
                "tags": {},
                "difficulties": {1: 0, 2: 0, 3: 0, 4: 0, 5: 0},
                "error_counts": [],
                "daily_added": {}
            }

            # 按日期统计新增题目
            date_format = "%Y-%m-%d"
            if start_date:
                current_date = start_date
                while current_date <= now:
                    stats_data["daily_added"][current_date.strftime(date_format)] = 0
                    current_date += timedelta(days=1)
            else:
                # 如果没有时间范围限制，统计所有日期的数据
                min_date = min(q.created_at for q in questions)
                max_date = max(q.created_at for q in questions)
                current_date = min_date
                while current_date <= max_date:
                    stats_data["daily_added"][current_date.strftime(date_format)] = 0
                    current_date += timedelta(days=1)

            for q in questions:
                # 知识点统计
                for tag in q.tags:
                    stats_data["tags"][tag] = stats_data["tags"].get(tag, 0) + q.error_count

                # 难度统计
                diff = int(round(q.difficulty))
                stats_data["difficulties"][diff] += 1

                # 错误次数
                stats_data["error_counts"].append(q.error_count)

                # 每日新增
                date_str = q.created_at.strftime(date_format)
                if date_str in stats_data["daily_added"]:
                    stats_data["daily_added"][date_str] += 1

            # 创建图表
            fig = plt.figure(figsize=(12, 8))
            fig.suptitle(f"错题统计分析 ({'全部数据' if not start_date else f'最近{time_range}'})")

            # 知识点分布 (饼图)
            ax1 = plt.subplot2grid((2, 2), (0, 0))
            if stats_data["tags"]:
                tags_sorted = dict(sorted(stats_data["tags"].items(), key=lambda x: x[1], reverse=True)[:8])
                ax1.pie(tags_sorted.values(), labels=tags_sorted.keys(), autopct="%1.1f%%")
                ax1.set_title("知识点分布 (按错误次数)")
            else:
                ax1.text(0.5, 0.5, "无知识点数据", ha="center", va="center")
                ax1.set_title("知识点分布")

            # 难度分布 (柱状图)
            ax2 = plt.subplot2grid((2, 2), (0, 1))
            ax2.bar(
                stats_data["difficulties"].keys(),
                stats_data["difficulties"].values(),
                color='skyblue'
            )
            ax2.set_title("难度分布")
            ax2.set_xlabel("难度等级")
            ax2.set_ylabel("题目数量")
            ax2.set_xticks([1, 2, 3, 4, 5])

            # 错误次数分布 (直方图)
            ax3 = plt.subplot2grid((2, 2), (1, 0))
            if stats_data["error_counts"]:
                ax3.hist(stats_data["error_counts"], bins=10, color='green', alpha=0.7)
                ax3.set_title("错误次数分布")
                ax3.set_xlabel("错误次数")
                ax3.set_ylabel("题目数量")
            else:
                ax3.text(0.5, 0.5, "无错误次数数据", ha="center", va="center")
                ax3.set_title("错误次数分布")

            # 每日新增题目 (折线图)
            ax4 = plt.subplot2grid((2, 2), (1, 1))
            if stats_data["daily_added"]:
                dates = sorted(stats_data["daily_added"].keys())
                counts = [stats_data["daily_added"][d] for d in dates]
                ax4.plot(dates, counts, marker='o', linestyle='-', color='orange')
                ax4.set_title("每日新增题目")
                ax4.set_xlabel("日期")
                ax4.set_ylabel("新增数量")
                plt.setp(ax4.get_xticklabels(), rotation=45, ha="right")
            else:
                ax4.text(0.5, 0.5, "无每日新增数据", ha="center", va="center")
                ax4.set_title("每日新增题目")

            plt.tight_layout()

            # 嵌入到Tkinter
            canvas = FigureCanvasTkAgg(fig, self.stats_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

            # 保存图表引用，以便后续导出
            self.current_stats = fig

        except Exception as e:
            messagebox.showerror("错误", f"生成统计图表失败:\n{str(e)}")
        finally:
            session.close()

    def export_charts(self):
        """导出统计图表为图片"""
        if not self.current_stats:
            messagebox.showwarning("警告", "没有可导出的统计图表！")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG图片", "*.png"), ("所有文件", "*.*")],
            title="保存统计图表"
        )

        if file_path:
            try:
                self.current_stats.savefig(file_path, dpi=300, bbox_inches='tight')
                messagebox.showinfo("成功", f"图表已保存到:\n{file_path}")

                # 询问是否打开文件
                if messagebox.askyesno("打开文件", "是否现在打开导出的图表？"):
                    webbrowser.open(file_path)
            except Exception as e:
                messagebox.showerror("错误", f"导出图表失败:\n{str(e)}")

    # ================== 试卷生成方法 ==================
    def generate_paper(self):
        """生成试卷预览"""
        session = Session()
        try:
            # 获取所有题目并按优先级排序
            query = session.query(Question)

            # 应用知识点筛选
            tags_filter = self.tags_filter_var.get().strip()
            if tags_filter:
                tags = [t.strip() for t in tags_filter.split(",") if t.strip()]
                query = query.filter(Question.tags.op("&&")(tags))  # PostgreSQL数组包含操作

            questions = query.all()

            if not questions:
                messagebox.showwarning("警告", "没有找到符合条件的题目！")
                return

            # 计算优先级
            sorted_questions = sorted(questions, key=self.calculate_priority, reverse=True)

            # 按数量限制
            count = min(self.paper_count_var.get(), len(sorted_questions))
            selected_questions = sorted_questions[:count]

            # 显示预览
            self.paper_preview.delete(1.0, tk.END)
            self.paper_preview.insert(tk.END, f"{self.paper_title_var.get()}\n\n", "title")

            for i, q in enumerate(selected_questions, 1):
                self.paper_preview.insert(tk.END, f"{i}. {q.content}\n", "question")
                self.paper_preview.insert(tk.END, f"【知识点】{', '.join(q.tags)}\n", "meta")
                self.paper_preview.insert(tk.END, f"【难度】{'★' * int(round(q.difficulty))}\n\n", "meta")

            # 保存当前选择的题目用于导出
            self.current_paper = selected_questions
            self.status_bar.config(text=f"已生成包含 {count} 道题目的试卷")

        except Exception as e:
            messagebox.showerror("错误", f"生成试卷失败:\n{str(e)}")
        finally:
            session.close()

    def calculate_priority(self, question):
        """计算题目优先级（带遗忘曲线）"""
        days = (datetime.now() - question.last_tested).days
        forgetting_factor = 0.56 * (0.94 ** days) + 0.06
        return question.error_count * forgetting_factor + question.difficulty * 0.5

    def export_to_word(self):
        """导出试卷为Word文档"""
        if not hasattr(self, 'current_paper') or not self.current_paper:
            messagebox.showwarning("警告", "请先生成试卷！")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            title="保存试卷"
        )

        if not file_path:
            return

        try:
            doc = Document()

            # 设置文档样式
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)

            # 添加标题
            title = doc.add_heading(self.paper_title_var.get(), level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加题目
            for i, q in enumerate(self.current_paper, 1):
                # 题目内容
                doc.add_paragraph(f"{i}. {q.content}")

                # 知识点和难度
                meta = doc.add_paragraph()
                meta.add_run("【知识点】").bold = True
                meta.add_run(f"{', '.join(q.tags)}  ")
                meta.add_run("【难度】").bold = True
                meta.add_run(f"{'★' * int(round(q.difficulty))}")

                # 添加空白行
                doc.add_paragraph()

            doc.save(file_path)
            messagebox.showinfo("成功", f"试卷已导出到:\n{file_path}")

            # 询问是否打开文件
            if messagebox.askyesno("打开文件", "是否现在打开导出的试卷？"):
                webbrowser.open(file_path)

        except ImportError:
            messagebox.showerror("错误", "请先安装python-docx库:\npip install python-docx")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败:\n{str(e)}")

    def export_to_pdf(self):
        """导出试卷为PDF"""
        messagebox.showinfo(
            "PDF导出",
            "建议先导出为Word文档，然后使用Word转换为PDF。\n"
            "或者安装pdfkit库: pip install pdfkit\n"
            "并安装wkhtmltopdf: https://wkhtmltopdf.org/"
        )

    # ================== 其他方法 ==================
    def on_close(self):
        """关闭窗口前的清理工作"""
        if messagebox.askokcancel("退出", "确定要退出智能错题系统吗？"):
            self.destroy()


# ================== 程序入口 ==================
if __name__ == "__main__":
    app = ExamToolApp()
    app.mainloop()