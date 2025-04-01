#阶段一
#数据存储与录入
#导入基本库
# 阶段一
# 数据存储与录入
# 导入基本库
import json # 导入json库，其作用为将Python数据结构转换为JSON格式数据，方便存储与传输
from datetime import datetime# 导入datetime库，用于记录错题最后测试时间
from docx import Document # 导入docx库，用于生成试卷与答案文档
import os # 导入os库，用于文件操作

DB_FILE = "questions.json"# 错题库文件名

# 初始化错题库
def init_database():
    if not os.path.exists(DB_FILE): #检查数据库文件是否存在
        with open(DB_FILE, "w") as f: #如果不存在创建一个新文件  with open(),"w"表示写入模式
            json.dump({"questions": []}, f)#将Python对象序列化为JSON格式 参数：数据对象“{"questions":[]}初始化空数组 文件对象：f


def add_question():
    print("\n=== 添加错题 ===")
    content = input("题目内容：")
    answer = input("正确答案：")
    tags = input("知识点标签（用逗号分隔）：").split(',')
#try-except块用于捕获异常，如文件打开失败等
    try:
        with open(DB_FILE, "r+") as f: #以读写模式打开文件
            data = json.load(f) #使用json.load方法将JSON格式数据转换为Python对象
            data["questions"].append({
                "id": len(data["questions"]) + 1,
                "content": content,
                "answer": answer,
                "tags": [tag.strip() for tag in tags],
                "error_count": 1,
                "last_tested": datetime.now().strftime("%Y-%m-%d")
            })
            f.seek(0) #将文件指针移动到开头，以便覆盖写入新的数据
            json.dump(data, f, indent=2) #将更新后的数据写回文件，indent=2使得JSON文件格式化，便于阅读。
        print("✅ 错题添加成功！")
    except Exception as e:
        print(f"❌ 保存失败：{e}")


def calculate_priority(q):
    """计算题目优先级（含时间衰减因子）"""
    last_tested = datetime.strptime(q["last_tested"], "%Y-%m-%d") #由于last_tested存储的是字符串，所以需要先转换为datetime格式
    days_passed = (datetime.now() - last_tested).days
    return q["error_count"] * (0.9 ** days_passed)  # 时间衰减公式


def generate_paper(num=10):
    try:
        with open(DB_FILE) as f:
            questions = json.load(f)["questions"] # 读取JSON文件中的questions数组
            sorted_questions = sorted(questions, key=calculate_priority, reverse=True) # sorted函数使用key参数 根据calculate_priority函数计算优先级排序，reverse参数控制降序排列
            return sorted_questions[:num] # 取前num个
    except FileNotFoundError:
        print("❌ 未找到错题库，请先添加错题！")
        return []


def mark_learned_questions(paper_questions):
    """标记已掌握的题目"""
    if not paper_questions:
        return

    print("\n=== 标记已掌握题目 ===")
    print("请输入已掌握的题号（如：1,3），按回车跳过：")
    learned_input = input().strip()

    if not learned_input:
        return

    try:
        learned_nums = [int(num) for num in learned_input.split(',')]
        learned_indices = [num - 1 for num in learned_nums]  # 转为索引
    except ValueError:
        print("❌ 输入格式错误，请使用逗号分隔的数字！")
        return

    try:
        with open(DB_FILE, "r+") as f:
            data = json.load(f)
            db_questions = data["questions"]

            # 构建ID到题目的映射
            id_to_question = {q["id"]: q for q in db_questions} #使用字典推导式。这行代码遍历列表中的每个题目q，并将其ID作为键，q作为值，构建一个字典。

#遍历learned_indices中的每个索引值。
            for idx in learned_indices:
                if 0 <= idx < len(paper_questions):
                    paper_q = paper_questions[idx]
                    q_id = paper_q["id"]
                    if q_id in id_to_question:
                        db_q = id_to_question[q_id]
                        # 减少错误次数（不低于0）
                        db_q["error_count"] = max(db_q["error_count"] - 1, 0)
                        # 更新最后测试时间
                        db_q["last_tested"] = datetime.now().strftime("%Y-%m-%d")

            f.seek(0)
            json.dump(data, f, indent=2)
            print(f"✅ 已更新 {len(learned_indices)} 道题目的掌握状态！")
    except Exception as e:
        print(f"❌ 更新失败：{e}")


def export_to_word(questions):
    if not questions:
        return

    # 生成试卷
    doc = Document()
    doc.add_heading('错题强化试卷', 0)
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q['content']}")
        doc.add_paragraph(f"【知识点】{', '.join(q['tags'])}\n")
    doc.save("试卷.docx")

    # 生成答案
    ans_doc = Document()
    ans_doc.add_heading('参考答案', 0)
    for i, q in enumerate(questions, 1):
        ans_doc.add_paragraph(f"{i}. {q['answer']}")
    ans_doc.save("答案.docx")

    print("✅ 已生成：试卷.docx 和 答案.docx")


def main():
    init_database()
    while True:
        print("\n=== 智能错题系统 ===")
        print("1. 添加错题")
        print("2. 生成试卷+答案")
        print("3. 退出")
        choice = input("请选择操作：")

        if choice == "1":
            add_question()
        elif choice == "2":
            num = int(input("请输入题目数量："))
            paper = generate_paper(num)
            export_to_word(paper)
            if paper:
                mark_learned_questions(paper)
        elif choice == "3":
            print("再见！")
            break
        else:
            print("无效输入，请重新选择！")


if __name__ == "__main__":
    main()
