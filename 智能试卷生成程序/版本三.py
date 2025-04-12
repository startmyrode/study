import os
import shutil
import tkinter as tk
import urllib.parse
import webbrowser
from collections import defaultdict
from datetime import datetime, timedelta
from tkinter import ttk, messagebox, filedialog, simpledialog
import matplotlib

matplotlib.use("TkAgg")  # 必须放在其他matplotlib导入之前
import matplotlib.pyplot as plt
from PIL import Image, ImageTk
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from sqlalchemy import (
    create_engine,
    Column,
    Integer,
    String,
    JSON,
    DateTime,
    Float,
    ForeignKey,
)
from sqlalchemy.orm import declarative_base, sessionmaker, relationship


# 配置数据库连接
def get_db_connection():
    # 请修改为你的实际数据库配置
    username = "CZQ"
    password = "1357810la"  # 替换为你的密码
    server = "devil"
    database = "exam_db"

    # 对密码中的特殊字符进行编码
    encoded_password = urllib.parse.quote_plus(password)

    # 使用ODBC Driver 17 for SQL Server
    return f"mssql+pyodbc://{username}:{encoded_password}@{server}/{database}?driver=ODBC+Driver+17+for+SQL+Server"


DATABASE_URL = get_db_connection()  # 获取数据库连接字符串
Base = declarative_base()  # 创建ORM模型基类


class Subject(Base):
    __tablename__ = "subjects"
    id = Column(Integer, primary_key=True, autoincrement=True)
    name = Column(String(50), unique=True, nullable=False)
    questions = relationship("Question", back_populates="subject")


class Question(Base):
    __tablename__ = "questions"
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(String(500))
    image_path = Column(String(200))
    answer = Column(String(500))
    tags = Column(JSON)
    error_count = Column(Integer, default=1)
    difficulty = Column(Float, default=3.0)
    last_tested = Column(DateTime, default=datetime.now)
    subject_id = Column(Integer, ForeignKey("subjects.id"))
    subject = relationship("Subject", back_populates="questions")


# 初始化数据库
engine = create_engine(DATABASE_URL, pool_pre_ping=True)
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)

# 图片存储目录
IMAGE_DIR = "question_images"
os.makedirs(IMAGE_DIR, exist_ok=True)


# ================== 主应用程序 ==================
class ExamToolApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("智能错题系统 v4.0 (支持图片+分科目)")
        self.geometry("1200x800")
        self.protocol("WM_DELETE_WINDOW", self.on_close)

        # 初始化变量
        self.current_paper = []
        self.current_stats = None
        self.current_subject = None
        self.editing_id = None
        self.image_path = None
        self.temp_images = []

        # 创建界面
        self.create_widgets()

        # 加载初始数据
        self.load_subjects()

    def create_widgets(self):
        """创建所有界面组件"""
        # 主菜单
        menubar = tk.Menu(self)
        self.config(menu=menubar)

        # 科目菜单
        subject_menu = tk.Menu(menubar, tearoff=0)
        subject_menu.add_command(label="添加科目", command=self.add_subject)
        subject_menu.add_command(label="删除科目", command=self.delete_subject)
        menubar.add_cascade(label="科目管理", menu=subject_menu)

        # 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="使用说明", command=self.show_help)
        menubar.add_cascade(label="帮助", menu=help_menu)

        # 主布局
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        # 标签页1：错题管理
        self.tab1 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text="错题管理")
        self.create_question_tab()

        # 标签页2：统计分析
        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab2, text="统计分析")
        self.create_stats_tab()

        # 标签页3：生成试卷
        self.tab3 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab3, text="生成试卷")
        self.create_paper_tab()

        # 状态栏
        self.status_bar = ttk.Label(
            self, text="就绪 | 当前科目: 未选择", relief=tk.SUNKEN
        )
        self.status_bar.pack(fill=tk.X)

    def create_question_tab(self):
        """创建错题管理标签页"""
        # 科目选择框
        subject_frame = ttk.Frame(self.tab1)
        subject_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(subject_frame, text="当前科目:").pack(side=tk.LEFT)
        self.subject_combo = ttk.Combobox(subject_frame, state="readonly")
        self.subject_combo.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        self.subject_combo.bind("<<ComboboxSelected>>", self.on_subject_selected)

        ttk.Button(subject_frame, text="刷新科目", command=self.load_subjects).pack(
            side=tk.LEFT, padx=5
        )

        # 输入表单框架
        form_frame = ttk.LabelFrame(self.tab1, text="题目内容")
        form_frame.pack(fill=tk.X, padx=10, pady=5)

        # 图片上传区域
        img_frame = ttk.Frame(form_frame)
        img_frame.grid(row=0, column=0, sticky="ns", padx=5, pady=5)

        self.img_label = ttk.Label(img_frame, text="[无题目图片]")
        self.img_label.pack()

        ttk.Button(img_frame, text="上传图片", command=self.upload_image).pack(pady=5)
        ttk.Button(img_frame, text="清除图片", command=self.clear_image).pack()
        ttk.Button(img_frame, text="查看原图", command=self.view_image).pack(pady=5)

        # 文本输入区域
        text_frame = ttk.Frame(form_frame)
        text_frame.grid(row=0, column=1, sticky="nsew")

        # 题目内容
        ttk.Label(text_frame, text="题目内容:").grid(
            row=0, column=0, sticky="w", pady=2
        )
        self.content_entry = tk.Text(text_frame, height=5, width=60, wrap=tk.WORD)
        self.content_entry.grid(row=0, column=1, columnspan=3, padx=5, sticky="ew")

        # 正确答案
        ttk.Label(text_frame, text="正确答案:").grid(
            row=1, column=0, sticky="w", pady=2
        )
        self.answer_entry = tk.Text(text_frame, height=3, width=60, wrap=tk.WORD)
        self.answer_entry.grid(row=1, column=1, columnspan=3, padx=5, sticky="ew")

        # 知识点标签
        ttk.Label(text_frame, text="知识点标签:").grid(
            row=2, column=0, sticky="w", pady=2
        )
        self.tags_entry = ttk.Entry(text_frame, width=60)
        self.tags_entry.grid(row=2, column=1, columnspan=3, padx=5, sticky="ew")
        ttk.Label(text_frame, text="(多个标签用逗号分隔)").grid(
            row=3, column=1, sticky="w"
        )

        # 难度和错误次数
        ttk.Label(text_frame, text="难度:").grid(row=4, column=0, sticky="w", pady=2)
        self.difficulty_var = tk.IntVar(value=3)
        ttk.Combobox(
            text_frame,
            textvariable=self.difficulty_var,
            values=[1, 2, 3, 4, 5],
            width=5,
        ).grid(row=4, column=1, sticky="w")

        ttk.Label(text_frame, text="错误次数:").grid(
            row=4, column=2, sticky="w", padx=10
        )
        self.error_count_var = tk.IntVar(value=1)
        ttk.Spinbox(
            text_frame, textvariable=self.error_count_var, from_=1, to=100, width=5
        ).grid(row=4, column=3, sticky="w")

        # 按钮组
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=1, column=0, columnspan=2, pady=10, sticky="e")

        ttk.Button(button_frame, text="清空", command=self.clear_form).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(button_frame, text="保存", command=self.save_question).pack(
            side=tk.LEFT, padx=5
        )

        # 错题列表框架
        list_frame = ttk.LabelFrame(self.tab1, text="错题列表")
        list_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

        # 列表工具栏
        toolbar = ttk.Frame(list_frame)
        toolbar.pack(fill=tk.X)

        ttk.Button(toolbar, text="刷新", command=self.load_questions).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Button(toolbar, text="标记为已掌握", command=self.mark_as_learned).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Button(toolbar, text="删除选中", command=self.delete_question).pack(
            side=tk.LEFT, padx=2
        )

        # 搜索框
        search_frame = ttk.Frame(toolbar)
        search_frame.pack(side=tk.RIGHT, padx=5)

        ttk.Label(search_frame, text="搜索:").pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=20)
        search_entry.pack(side=tk.LEFT, padx=2)
        search_entry.bind("<Return>", lambda e: self.load_questions())
        ttk.Button(search_frame, text="搜索", command=self.load_questions).pack(
            side=tk.LEFT
        )

        # 错题列表
        columns = (
            "id",
            "content",
            "tags",
            "error_count",
            "difficulty",
            "last_tested",
            "has_image",
        )
        self.tree = ttk.Treeview(
            list_frame, columns=columns, show="headings", selectmode="extended"
        )

        # 配置列
        self.tree.column("id", width=50, anchor="center")
        self.tree.column("content", width=300)
        self.tree.column("tags", width=150)
        self.tree.column("error_count", width=80, anchor="center")
        self.tree.column("difficulty", width=80, anchor="center")
        self.tree.column("last_tested", width=120, anchor="center")
        self.tree.column("has_image", width=80, anchor="center")

        # 设置列标题
        self.tree.heading("id", text="ID")
        self.tree.heading("content", text="题目内容")
        self.tree.heading("tags", text="知识点")
        self.tree.heading("error_count", text="错误次数")
        self.tree.heading("difficulty", text="难度")
        self.tree.heading("last_tested", text="最后测试")
        self.tree.heading("has_image", text="有无图片")

        # 添加滚动条
        scroll_y = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        scroll_y.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=scroll_y.set)

        scroll_x = ttk.Scrollbar(
            list_frame, orient="horizontal", command=self.tree.xview
        )
        scroll_x.pack(side="bottom", fill="x")
        self.tree.configure(xscrollcommand=scroll_x.set)

        self.tree.pack(fill=tk.BOTH, expand=True)

        # 绑定事件
        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.tree.bind("<Button-3>", self.show_context_menu)

        # 右键菜单
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="编辑", command=self.edit_question)
        self.context_menu.add_command(
            label="标记为已掌握", command=self.mark_as_learned
        )
        self.context_menu.add_command(
            label="查看图片", command=self.view_selected_image
        )
        self.context_menu.add_separator()
        self.context_menu.add_command(label="删除", command=self.delete_question)

    def create_stats_tab(self):
        """创建统计分析标签页"""
        self.stats_frame = ttk.Frame(self.tab2)
        self.stats_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 控制面板
        control_frame = ttk.Frame(self.tab2)
        control_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(control_frame, text="刷新统计", command=self.update_stats).pack(
            side=tk.LEFT
        )
        ttk.Button(control_frame, text="导出图表", command=self.export_charts).pack(
            side=tk.LEFT, padx=5
        )

        # 时间范围选择
        ttk.Label(control_frame, text="统计时间范围:").pack(side=tk.LEFT, padx=5)
        self.time_range_var = tk.StringVar(value="all")
        ttk.Radiobutton(
            control_frame,
            text="全部",
            variable=self.time_range_var,
            value="all",
            command=self.update_stats,
        ).pack(side=tk.LEFT)
        ttk.Radiobutton(
            control_frame,
            text="最近7天",
            variable=self.time_range_var,
            value="7days",
            command=self.update_stats,
        ).pack(side=tk.LEFT)
        ttk.Radiobutton(
            control_frame,
            text="最近30天",
            variable=self.time_range_var,
            value="30days",
            command=self.update_stats,
        ).pack(side=tk.LEFT)

        # 图表类型选择
        ttk.Label(control_frame, text="图表类型:").pack(side=tk.LEFT, padx=5)
        self.chart_type_var = tk.StringVar(value="combined")
        ttk.Radiobutton(
            control_frame,
            text="综合视图",
            variable=self.chart_type_var,
            value="combined",
            command=self.update_stats,
        ).pack(side=tk.LEFT)
        ttk.Radiobutton(
            control_frame,
            text="知识点分布",
            variable=self.chart_type_var,
            value="tags",
            command=self.update_stats,
        ).pack(side=tk.LEFT)
        ttk.Radiobutton(
            control_frame,
            text="错误趋势",
            variable=self.chart_type_var,
            value="trend",
            command=self.update_stats,
        ).pack(side=tk.LEFT)

    def create_paper_tab(self):
        """创建生成试卷标签页"""
        # 设置框架
        settings_frame = ttk.LabelFrame(self.tab3, text="试卷设置")
        settings_frame.pack(padx=10, pady=5, fill=tk.X)

        # 试卷标题
        ttk.Label(settings_frame, text="试卷标题:").grid(
            row=0, column=0, sticky="w", padx=5, pady=2
        )
        self.paper_title_var = tk.StringVar(value="错题强化试卷")
        ttk.Entry(settings_frame, textvariable=self.paper_title_var, width=40).grid(
            row=0, column=1, sticky="ew"
        )

        # 题目数量
        ttk.Label(settings_frame, text="题目数量:").grid(
            row=1, column=0, sticky="w", padx=5, pady=2
        )
        self.paper_count_var = tk.IntVar(value=10)
        ttk.Spinbox(
            settings_frame, textvariable=self.paper_count_var, from_=1, to=100, width=5
        ).grid(row=1, column=1, sticky="w")

        # 难度分布
        ttk.Label(settings_frame, text="难度分布:").grid(
            row=2, column=0, sticky="w", padx=5, pady=2
        )
        self.difficulty_dist_var = tk.StringVar(value="1:2:3:2:1")
        ttk.Entry(settings_frame, textvariable=self.difficulty_dist_var, width=15).grid(
            row=2, column=1, sticky="w"
        )
        ttk.Label(settings_frame, text="(格式:简单:较易:中等:较难:困难)").grid(
            row=2, column=2, sticky="w"
        )

        # 知识点筛选
        ttk.Label(settings_frame, text="知识点筛选:").grid(
            row=3, column=0, sticky="w", padx=5, pady=2
        )
        self.tags_filter_var = tk.StringVar()
        ttk.Entry(settings_frame, textvariable=self.tags_filter_var, width=40).grid(
            row=3, column=1, columnspan=2, sticky="ew"
        )

        # 生成按钮
        button_frame = ttk.Frame(settings_frame)
        button_frame.grid(row=4, column=1, columnspan=2, pady=5, sticky="e")

        ttk.Button(button_frame, text="生成试卷", command=self.generate_paper).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(button_frame, text="导出Word", command=self.export_to_word).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(button_frame, text="导出PDF", command=self.export_to_pdf).pack(
            side=tk.LEFT, padx=5
        )

        # 试卷预览
        preview_frame = ttk.LabelFrame(self.tab3, text="试卷预览")
        preview_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

        self.paper_preview = tk.Text(
            preview_frame, wrap=tk.WORD, font=("宋体", 10), padx=5, pady=5
        )

        scroll_y = ttk.Scrollbar(
            preview_frame, orient="vertical", command=self.paper_preview.yview
        )
        scroll_y.pack(side="right", fill="y")
        self.paper_preview.configure(yscrollcommand=scroll_y.set)

        scroll_x = ttk.Scrollbar(
            preview_frame, orient="horizontal", command=self.paper_preview.xview
        )
        scroll_x.pack(side="bottom", fill="x")
        self.paper_preview.configure(xscrollcommand=scroll_x.set)

        self.paper_preview.pack(fill=tk.BOTH, expand=True)

    # ================== 数据库操作方法 ==================
    def load_subjects(self):
        """加载所有科目到下拉框"""
        session = Session()
        try:
            subjects = session.query(Subject).order_by(Subject.name).all()
            self.subject_combo["values"] = [s.name for s in subjects]

            if subjects and not self.current_subject:
                self.current_subject = subjects[0]
                self.subject_combo.set(self.current_subject.name)
                self.status_bar.config(
                    text=f"就绪 | 当前科目: {self.current_subject.name}"
                )
                self.load_questions()
        except Exception as e:
            messagebox.showerror("错误", f"加载科目失败:\n{str(e)}")
        finally:
            session.close()

    def add_subject(self):
        """添加新科目"""
        subject_name = simpledialog.askstring("添加科目", "请输入科目名称:")
        if subject_name:
            session = Session()
            try:
                existing = session.query(Subject).filter_by(name=subject_name).first()
                if existing:
                    messagebox.showwarning("警告", f"科目'{subject_name}'已存在！")
                    return

                new_subject = Subject(name=subject_name)
                session.add(new_subject)
                session.commit()
                messagebox.showinfo("成功", f"科目'{subject_name}'已添加")
                self.load_subjects()
            except Exception as e:
                session.rollback()
                messagebox.showerror("错误", f"添加失败:\n{str(e)}")
            finally:
                session.close()

    def delete_subject(self):
        """删除当前科目"""
        if not self.current_subject:
            messagebox.showwarning("警告", "请先选择要删除的科目！")
            return

        if not messagebox.askyesno(
            "确认",
            f"确定要删除科目'{self.current_subject.name}'吗？\n这将删除该科目下的所有错题记录！",
        ):
            return

        session = Session()
        try:
            # 先删除相关错题图片
            questions = (
                session.query(Question)
                .filter_by(subject_id=self.current_subject.id)
                .all()
            )
            for q in questions:
                if q.image_path and os.path.exists(q.image_path):
                    try:
                        os.remove(q.image_path)
                    except:
                        pass

            # 删除科目及其关联题目
            session.delete(self.current_subject)
            session.commit()

            messagebox.showinfo("成功", f"科目'{self.current_subject.name}'已删除")
            self.current_subject = None
            self.subject_combo.set("")
            self.status_bar.config(text="就绪 | 当前科目: 未选择")
            self.load_subjects()
            self.clear_form()
            self.tree.delete(*self.tree.get_children())
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"删除失败:\n{str(e)}")
        finally:
            session.close()

    def on_subject_selected(self, event):
        """科目选择变化事件"""
        subject_name = self.subject_combo.get()
        session = Session()
        try:
            self.current_subject = (
                session.query(Subject).filter_by(name=subject_name).first()
            )
            if self.current_subject:
                self.status_bar.config(
                    text=f"就绪 | 当前科目: {self.current_subject.name}"
                )
                self.load_questions()
        except Exception as e:
            messagebox.showerror("错误", f"加载科目数据失败:\n{str(e)}")
        finally:
            session.close()

    def load_questions(self):
        """加载当前科目的错题"""
        self.tree.delete(*self.tree.get_children())

        if not self.current_subject:
            return

        session = Session()
        try:
            query = session.query(Question).filter_by(
                subject_id=self.current_subject.id
            )

            # 应用搜索条件
            search_text = self.search_var.get().strip()
            if search_text:
                query = query.filter(
                    Question.content.like(f"%{search_text}%")
                    | Question.answer.like(f"%{search_text}%")
                    | Question.tags.op("&&")([search_text])  # PostgreSQL数组包含操作
                )

            # 按优先级排序
            questions = sorted(query.all(), key=self.calculate_priority, reverse=True)

            for q in questions:
                self.tree.insert(
                    "",
                    tk.END,
                    values=(
                        q.id,
                        q.content[:100] + "..." if len(q.content) > 100 else q.content,
                        ", ".join(q.tags) if q.tags else "",
                        q.error_count,
                        q.difficulty,
                        q.last_tested.strftime("%Y-%m-%d"),
                        "有" if q.image_path else "无",
                    ),
                )

            self.status_bar.config(text=f"加载完成，共 {len(questions)} 道错题")
        except Exception as e:
            messagebox.showerror("错误", f"加载数据失败:\n{str(e)}")
        finally:
            session.close()

    def save_question(self):
        """保存题目到数据库"""
        if not self.current_subject:
            messagebox.showerror("错误", "请先选择科目！")
            return

        content = self.content_entry.get("1.0", tk.END).strip()
        answer = self.answer_entry.get("1.0", tk.END).strip()
        tags = [t.strip() for t in self.tags_entry.get().split(",") if t.strip()]

        if not content:
            messagebox.showwarning("警告", "题目内容不能为空！")
            return

        session = Session()
        try:
            # 处理图片路径
            img_path = None
            if self.image_path:
                # 如果是新上传的图片（临时路径），则移动到正式目录
                if not os.path.dirname(self.image_path) == os.path.abspath(IMAGE_DIR):
                    filename = f"img_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.path.basename(self.image_path)}"
                    dest_path = os.path.join(IMAGE_DIR, filename)
                    shutil.copy(self.image_path, dest_path)
                    img_path = dest_path
                else:
                    img_path = self.image_path

            if self.editing_id:  # 编辑模式
                question = session.query(Question).get(self.editing_id)

                # 删除旧图片（如果有变化）
                if question.image_path and question.image_path != img_path:
                    try:
                        os.remove(question.image_path)
                    except:
                        pass

                question.content = content
                question.image_path = img_path
                question.answer = answer
                question.tags = tags
                question.difficulty = self.difficulty_var.get()
                question.error_count = self.error_count_var.get()
                question.last_tested = datetime.now()
            else:  # 新增模式
                new_question = Question(
                    content=content,
                    image_path=img_path,
                    answer=answer,
                    tags=tags,
                    difficulty=self.difficulty_var.get(),
                    error_count=self.error_count_var.get(),
                    last_tested=datetime.now(),
                    subject_id=self.current_subject.id,
                )
                session.add(new_question)

            session.commit()
            messagebox.showinfo("成功", "题目已保存！")

            self.clear_form()
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"保存失败:\n{str(e)}")
        finally:
            session.close()

    def clear_form(self):
        """清空输入表单"""
        self.content_entry.delete("1.0", tk.END)
        self.answer_entry.delete("1.0", tk.END)
        self.tags_entry.delete(0, tk.END)
        self.difficulty_var.set(3)
        self.error_count_var.set(1)
        self.clear_image()
        self.editing_id = None

        # 清除树形视图的选择
        for item in self.tree.selection():
            self.tree.selection_remove(item)

    def upload_image(self):
        """上传题目图片"""
        file_path = filedialog.askopenfilename(
            filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif")]
        )
        if file_path:
            self.image_path = file_path
            self.show_image_thumbnail()

    def show_image_thumbnail(self):
        """显示图片缩略图"""
        if self.image_path and os.path.exists(self.image_path):
            try:
                img = Image.open(self.image_path)
                img.thumbnail((200, 200))
                photo = ImageTk.PhotoImage(img)

                self.img_label.config(image=photo)
                self.img_label.image = photo  # 保持引用
                self.img_label.text = ""
            except Exception as e:
                messagebox.showerror("错误", f"无法加载图片:\n{str(e)}")
                self.clear_image()

    def clear_image(self):
        """清除当前图片"""
        self.image_path = None
        self.img_label.config(image="", text="[无题目图片]")

    def view_image(self):
        """查看当前图片原图"""
        if self.image_path and os.path.exists(self.image_path):
            try:
                img = Image.open(self.image_path)
                img.show()
            except Exception as e:
                messagebox.showerror("错误", f"无法打开图片:\n{str(e)}")
        else:
            messagebox.showinfo("提示", "当前没有可查看的图片")

    def view_selected_image(self):
        """查看选中题目的图片"""
        selected_item = self.tree.focus()
        if not selected_item:
            return

        item_id = self.tree.item(selected_item)["values"][0]
        session = Session()
        try:
            question = session.query(Question).get(item_id)
            if question and question.image_path and os.path.exists(question.image_path):
                img = Image.open(question.image_path)
                img.show()
            else:
                messagebox.showinfo("提示", "该题目没有图片或图片已丢失")
        except Exception as e:
            messagebox.showerror("错误", f"无法打开图片:\n{str(e)}")
        finally:
            session.close()

    def on_tree_double_click(self, event):
        """双击错题列表项加载到编辑表单"""
        self.edit_question()

    def edit_question(self):
        """编辑选中的错题"""
        selected_item = self.tree.focus()
        if not selected_item:
            return

        item_id = self.tree.item(selected_item)["values"][0]
        session = Session()
        try:
            question = session.query(Question).get(item_id)
            if question:
                self.clear_form()
                self.editing_id = question.id

                self.content_entry.insert("1.0", question.content)
                self.answer_entry.insert("1.0", question.answer)
                self.tags_entry.insert(
                    0, ", ".join(question.tags) if question.tags else ""
                )
                self.difficulty_var.set(question.difficulty)
                self.error_count_var.set(question.error_count)

                # 加载图片
                if question.image_path and os.path.exists(question.image_path):
                    self.image_path = question.image_path
                    self.show_image_thumbnail()

                self.status_bar.config(text=f"正在编辑题目 ID: {item_id}")
        except Exception as e:
            messagebox.showerror("错误", f"加载题目失败:\n{str(e)}")
        finally:
            session.close()

    def mark_as_learned(self):
        """标记选中的题目为已掌握"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择题目！")
            return

        if not messagebox.askyesno(
            "确认", f"确定要标记 {len(selected_items)} 道题为已掌握吗？"
        ):
            return

        session = Session()
        try:
            count = 0
            for item in selected_items:
                item_id = self.tree.item(item)["values"][0]
                question = session.query(Question).get(item_id)
                if question:
                    question.error_count = max(0, question.error_count - 1)
                    question.last_tested = datetime.now()
                    count += 1

            session.commit()
            messagebox.showinfo("成功", f"已标记 {count} 道题为已掌握")
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"标记失败:\n{str(e)}")
        finally:
            session.close()

    def delete_question(self):
        """删除选中的题目"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择题目！")
            return

        if not messagebox.askyesno(
            "确认", f"确定要删除 {len(selected_items)} 道题吗？此操作不可恢复！"
        ):
            return

        session = Session()
        try:
            count = 0
            for item in selected_items:
                item_id = self.tree.item(item)["values"][0]
                question = session.query(Question).get(item_id)
                if question:
                    # 删除关联的图片文件
                    if question.image_path and os.path.exists(question.image_path):
                        try:
                            os.remove(question.image_path)
                        except:
                            pass

                    session.delete(question)
                    count += 1

            session.commit()
            messagebox.showinfo("成功", f"已删除 {count} 道题")
            self.load_questions()
            self.update_stats()
        except Exception as e:
            session.rollback()
            messagebox.showerror("错误", f"删除失败:\n{str(e)}")
        finally:
            session.close()

    def show_context_menu(self, event):
        """显示右键菜单"""
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

        # ================== 统计分析方法 ==================

    def update_stats(self):
        """更新统计图表"""
        # 设置中文字体
        plt.rcParams["font.sans-serif"] = [
            "SimHei",
            "Microsoft YaHei",
            "WenQuanYi Zen Hei",
        ]  # 多个备选字体
        plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题
        # 清除旧内容
        for widget in self.stats_frame.winfo_children():
            widget.destroy()

        if not self.current_subject:
            ttk.Label(self.stats_frame, text="请先选择科目").pack()
            return

        session = Session()
        try:
            # 确定时间范围
            time_range = self.time_range_var.get()
            now = datetime.now()

            if time_range == "7days":
                start_date = now - timedelta(days=7)
            elif time_range == "30days":
                start_date = now - timedelta(days=30)
            else:
                start_date = None  # 全部数据

            # 获取数据
            query = session.query(Question).filter_by(
                subject_id=self.current_subject.id
            )
            if start_date:
                query = query.filter(Question.last_tested >= start_date)

            questions = query.all()

            if not questions:
                ttk.Label(self.stats_frame, text="当前科目没有错题数据").pack()
                return

            # 准备统计数据
            stats_data = {
                "tags": defaultdict(int),
                "difficulties": {1: 0, 2: 0, 3: 0, 4: 0, 5: 0},
                "error_counts": [],
                "daily_errors": defaultdict(int),
            }

            # 按日期统计
            date_format = "%Y-%m-%d"
            if start_date:
                current_date = start_date
                while current_date <= now:
                    date_str = current_date.strftime(date_format)
                    stats_data["daily_errors"][date_str] = 0
                    current_date += timedelta(days=1)

            for q in questions:
                # 知识点统计
                if q.tags:
                    for tag in q.tags:
                        stats_data["tags"][tag] += q.error_count

                # 难度统计
                diff = int(round(q.difficulty))
                stats_data["difficulties"][diff] += 1

                # 错误次数
                stats_data["error_counts"].append(q.error_count)

                # 每日错误次数
                date_str = q.last_tested.strftime(date_format)
                stats_data["daily_errors"][date_str] += q.error_count

            # 创建图表
            chart_type = self.chart_type_var.get()
            fig = plt.figure(figsize=(10, 6), tight_layout=True)
            fig.patch.set_facecolor("#F0F0F0")

            if chart_type == "combined":
                # 综合视图
                fig.suptitle(
                    f"【{self.current_subject.name}】错题综合分析\n（{time_range if start_date else '全部数据'}）",
                    fontsize=12,
                    y=1.05,
                )

                # 知识点分布
                ax1 = plt.subplot2grid((2, 2), (0, 0))
                if stats_data["tags"]:
                    tags_sorted = sorted(
                        stats_data["tags"].items(), key=lambda x: x[1], reverse=True
                    )[:8]
                    labels = [t[0] for t in tags_sorted]
                    sizes = [t[1] for t in tags_sorted]
                    ax1.pie(sizes, labels=labels, autopct="%1.1f%%", startangle=90)
                    ax1.set_title("知识点分布TOP8")
                else:
                    ax1.text(0.5, 0.5, "无知识点数据", ha="center", va="center")

                # 难度分布
                ax2 = plt.subplot2grid((2, 2), (0, 1))
                diffs = sorted(stats_data["difficulties"].items())
                ax2.bar([d[0] for d in diffs], [d[1] for d in diffs])
                ax2.set_xticks([1, 2, 3, 4, 5])
                ax2.set_xlabel("难度等级")
                ax2.set_ylabel("题目数量")
                ax2.set_title("难度分布")

                # 错误趋势
                ax3 = plt.subplot2grid((2, 2), (1, 0), colspan=2)
                dates = sorted(stats_data["daily_errors"].keys())
                values = [stats_data["daily_errors"][d] for d in dates]
                ax3.plot(dates, values, marker="o")
                step = max(1, len(dates) // 5)
                ax3.set_xticks(dates[::step])
                plt.setp(ax3.get_xticklabels(), rotation=45, ha="right")
                ax3.set_title("每日错误次数趋势")

            elif chart_type == "tags":
                # 知识点分布（柱状图）
                ax = fig.add_subplot(111)
                if stats_data["tags"]:
                    tags_sorted = sorted(
                        stats_data["tags"].items(), key=lambda x: x[1], reverse=True
                    )[:10]
                    ax.barh([t[0] for t in tags_sorted], [t[1] for t in tags_sorted])
                    ax.set_title(
                        f"知识点错误次数TOP10（{time_range if start_date else '全部数据'}）"
                    )
                    ax.invert_yaxis()
                    plt.subplots_adjust(left=0.3)
                else:
                    ax.text(0.5, 0.5, "无知识点数据", ha="center", va="center")

            elif chart_type == "trend":
                # 错误趋势（双轴图）
                ax = fig.add_subplot(111)
                dates = sorted(stats_data["daily_errors"].keys())
                values = [stats_data["daily_errors"][d] for d in dates]

                # 错误次数折线图
                ax.plot(dates, values, "b-o", label="错误次数")
                ax.set_xticks(dates[:: len(dates) // 5])
                plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
                ax.set_ylabel("错误次数", color="b")
                ax.tick_params("y", colors="b")

                # 新增题目次轴
                ax2 = ax.twinx()
                ax2.plot(
                    dates,
                    [stats_data["daily_errors"][d] for d in dates],
                    "r--s",
                    label="新增题目",
                )
                ax2.set_ylabel("新增题目数", color="r")
                ax2.tick_params("y", colors="r")

                ax.set_title(
                    f"学习趋势分析（{time_range if start_date else '全部数据'}）"
                )
                fig.legend(loc="upper right", bbox_to_anchor=(0.9, 0.9))

            # 嵌入图表到界面
            canvas = FigureCanvasTkAgg(fig, master=self.stats_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        except Exception as e:
            messagebox.showerror("错误", f"生成图表失败:\n{str(e)}")

        finally:
            session.close()

    def export_charts(self):
        """导出当前图表为图片"""
        if not hasattr(self, "stats_frame") or not self.stats_frame.winfo_children():
            messagebox.showwarning("警告", "没有可导出的图表")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG图片", "*.png"), ("所有文件", "*.*")],
        )
        if file_path:
            try:
                # 获取当前显示的图表
                canvas = self.stats_frame.winfo_children()[0]
                figure = canvas.figure
                figure.savefig(file_path, dpi=300, bbox_inches="tight")
                messagebox.showinfo("成功", f"图表已保存到:\n{file_path}")
            except Exception as e:
                messagebox.showerror("错误", f"保存失败:\n{str(e)}")

    def generate_paper(self):
        """生成试卷"""
        if not self.current_subject:
            messagebox.showerror("错误", "请先选择科目！")
            return

        try:
            # 解析难度分布
            dist = [int(x) for x in self.difficulty_dist_var.get().split(":")]
            if len(dist) != 5 or sum(dist) == 0:
                raise ValueError("难度分布格式错误")

            total = self.paper_count_var.get()
            dist_ratio = [x / sum(dist) for x in dist]
            questions_by_diff = {i + 1: [] for i in range(5)}

            session = Session()
            try:
                # 获取符合条件的题目
                base_query = (
                    session.query(Question)
                    .filter_by(subject_id=self.current_subject.id)
                    .order_by(Question.last_tested)
                )

                # 应用标签筛选
                tags_filter = self.tags_filter_var.get().strip()
                if tags_filter:
                    tags = [t.strip() for t in tags_filter.split(",")]
                    base_query = base_query.filter(Question.tags.op("&&")(tags))

                # 按难度分类
                for diff in questions_by_diff:
                    query = base_query.filter(
                        Question.difficulty.between(diff - 0.5, diff + 0.5)
                    )
                    questions_by_diff[diff] = sorted(
                        query.all(), key=self.calculate_priority, reverse=True
                    )

                # 按比例选择题目
                self.current_paper = []
                for diff, count in enumerate(dist, 1):
                    select_num = min(
                        int(total * dist_ratio[diff - 1]),
                        len(questions_by_diff[diff]),
                    )
                    self.current_paper.extend(questions_by_diff[diff][:select_num])

                # 打乱顺序
                import random

                random.shuffle(self.current_paper)

                # 更新预览
                self.update_paper_preview()
                messagebox.showinfo(
                    "成功", f"已生成包含 {len(self.current_paper)} 道题的试卷"
                )

            finally:
                session.close()

        except Exception as e:
            messagebox.showerror("错误", f"生成试卷失败:\n{str(e)}")

    def update_paper_preview(self):
        """更新试卷预览"""
        self.paper_preview.delete(1.0, tk.END)
        text = f"【{self.paper_title_var.get()}】\n\n"
        text += f"科目：{self.current_subject.name}\n"
        text += f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"

        for i, q in enumerate(self.current_paper, 1):
            text += f"{i}. {q.content}\n"
            if q.image_path and os.path.exists(q.image_path):
                text += f"[图片] {os.path.basename(q.image_path)}\n"
            text += "\n"

        self.paper_preview.insert(tk.END, text)

    def export_to_word(self):
        """导出试卷到Word文档"""
        if not self.current_paper:
            messagebox.showwarning("警告", "请先生成试卷！")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
        )
        if not file_path:
            return

        try:
            doc = Document()
            # 设置文档样式
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(10.5)

            # 标题
            title = doc.add_paragraph(self.paper_title_var.get())
            title.style = doc.styles["Heading 1"]
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 基本信息
            doc.add_paragraph(f"科目：{self.current_subject.name}")
            doc.add_paragraph(
                f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            )

            # 添加题目
            for i, q in enumerate(self.current_paper, 1):
                # 题目内容
                p = doc.add_paragraph(f"{i}. {q.content}")

                # 插入图片
                if q.image_path and os.path.exists(q.image_path):
                    try:
                        doc.add_picture(q.image_path, width=Inches(4.0))
                    except:
                        doc.add_paragraph("[图片加载失败]")

                # 添加答案区域
                doc.add_paragraph("答案：" + "_" * 50 + "\n")

            doc.save(file_path)
            messagebox.showinfo("成功", f"试卷已保存到:\n{file_path}")
            # 自动打开文档
            webbrowser.open(file_path)
        except Exception as e:
            messagebox.showerror("错误", f"导出失败:\n{str(e)}")

    def export_to_pdf(self):
        """导出试卷为PDF（通过Word转换）"""
        if not self.current_paper:
            messagebox.showwarning("警告", "请先生成试卷！")
            return

        # 先保存Word文档
        word_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word文档", "*.docx")]
        )
        if not word_path:
            return

        self.export_to_word(word_path)  # 调用已有的Word导出方法

        try:
            # 转换为PDF
            pdf_path = os.path.splitext(word_path)[0] + ".pdf"

            import win32com.client

            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(word_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17是PDF格式代码
            doc.Close()
            word.Quit()

            messagebox.showinfo("成功", f"PDF已保存到:\n{pdf_path}")
            webbrowser.open(pdf_path)
        except Exception as e:
            messagebox.showerror("错误", f"PDF转换失败，请确认已安装Word\n{str(e)}")

        # 方法实现内容将放在这里

    def calculate_priority(self, question):
        """计算题目优先级（用于排序）"""
        days_since = (datetime.now() - question.last_tested).days
        priority = (
            question.error_count * 0.5
            + question.difficulty * 0.3
            + (1 / (days_since + 1)) * 0.2
        )
        return priority

    def show_help(self):
        """显示帮助文档"""
        help_text = """【使用说明】

        1. 科目管理：
           - 首次使用需通过菜单栏添加科目
           - 选择科目后才能进行相关操作

        2. 错题管理：
           - 上传题目时可添加图片（支持jpg/png等格式）
           - 双击题目进行编辑，右键菜单提供更多操作
           - 使用知识点标签进行分类管理

        3. 统计分析：
           - 支持按时间范围筛选数据
           - 提供综合视图、知识点分布、错误趋势三种图表
           - 点击"导出图表"保存为图片


        4. 生成试卷：
           - 设置难度分布（如1:2:3:2:1表示五个难度等级的题目比例）
           - 生成后可导出Word文档（自动包含题目图片）
           - PDF导出需要系统安装Word并配置打印功能

        5. 数据安全：
           - 所有数据存储在本地SQL Server数据库
           - 图片文件保存在程序目录下的question_images文件夹
           - 定期备份数据库文件（exam_db）以防止数据丢失
            """
        messagebox.showinfo("系统帮助", help_text)

    def on_close(self):
        """关闭窗口事件处理"""
        # 清理临时图片
        for img_path in self.temp_images:
            try:
                os.remove(img_path)
            except:
                pass
        self.destroy()


if __name__ == "__main__":
    app = ExamToolApp()
    app.mainloop()
    # 在应用启动时执行
